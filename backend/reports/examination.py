from docx import Document

from .models import Rule, Style


class RuleParams:

    def __init__(self, style: str) -> None:
        self.style = style
        self.rule_params: dict = {}
        for rule in Rule.objects.filter(style__slug=style):
            try:
                rule.value = int(rule.value)
            except ValueError:
                try:
                    rule.value = float(rule.value)
                except ValueError:
                    if rule.value in ['True', 'False']:
                        rule.value = rule.value == 'True'
            finally:
                self.rule_params[rule.param.slug] = (rule.param.name,
                                                     rule.value)

    def validate(self, data):
        errors = {}
        for par, value in self.rule_params.items():
            if data[par] != value[1]:
                errors[par] = {
                    'error_text': (
                        f"У вас {data[par]}, а должно быть "
                        f"{value[1]}."
                    ),
                    'name': value[0],
                }
        if self.style == 'Heading' and data['text'] != data['text'].upper():
            errors['upper'] = {
                'error_text': 'Заголовки должны быть в верхнем регистре.',
                'name': 'Верхний регистр',
            }
        return errors


def get_par_info(doc: Document) -> dict:
    detected_classes = {
        'Normal': RuleParams('Normal'),
        'Heading': RuleParams('Heading'),
        'List': RuleParams('List'),
        'ImageName': RuleParams('ImageName'),
        'TableName': RuleParams('TableName'),
    }
    style_info = {style.slug: style.name for style in Style.objects.all()}

    all_par_info = []

    paragraphs = doc.paragraphs

    for i, p in enumerate(paragraphs):
        par_info = dict()

        if p.text:
            style = p.style.name.split()[0]
            pr = p.runs[0]
            ps = p.style
            psbs = ps.base_style
            par_info['font_size'] = (
                int(pr.font.size.pt)
                if pr.font.size
                else int(ps.font.size.pt)
                if ps.font.size
                else int(psbs.font.size.pt)
                if psbs and psbs.font.size
                else None
            )
            par_info['font_name'] = (
                pr.font.name
                if pr.font.name
                else ps.font.name
                if ps.font.name
                else psbs.font.name
                if psbs and psbs.font.name
                else None
            )
            par_info['alignment'] = (
                p.alignment.name
                if p.alignment
                else ps.paragraph_format.alignment.name
                if ps.paragraph_format.alignment
                else psbs.paragraph_format.alignment.name
                if psbs and psbs.paragraph_format.alignment
                else 'LEFT'
            )
            par_info['line_spacing'] = (
                p.paragraph_format.line_spacing
                if p.paragraph_format.line_spacing
                else ps.paragraph_format.line_spacing
                if ps.paragraph_format.line_spacing
                else psbs.paragraph_format.line_spacing
                if psbs and psbs.paragraph_format.line_spacing
                else 1
            )
            par_info['first_line_indent'] = (
                float(f'{p.paragraph_format.first_line_indent.cm:.2f}')
                if p.paragraph_format.first_line_indent
                else float(f'{ps.paragraph_format.first_line_indent.cm:.2f}')
                if ps.paragraph_format.first_line_indent
                else float(f'{psbs.paragraph_format.first_line_indent.cm:.2f}')
                if psbs and psbs.paragraph_format.first_line_indent
                else 0
            )
            par_info['left_indent'] = (
                float(f'{p.paragraph_format.left_indent.cm:.2f}')
                if p.paragraph_format.left_indent
                else float(f'{ps.paragraph_format.left_indent.cm:.2f}')
                if ps.paragraph_format.left_indent
                else float(f'{psbs.paragraph_format.left_indent.cm:.2f}')
                if psbs and psbs.paragraph_format.left_indent
                else 0
            )
            par_info['right_indent'] = (
                float(f'{p.paragraph_format.right_indent.cm:.2f}')
                if p.paragraph_format.right_indent
                else float(f'{ps.paragraph_format.right_indent.cm:.2f}')
                if ps.paragraph_format.right_indent
                else float(f'{psbs.paragraph_format.right_indent.cm:.2f}')
                if psbs and psbs.paragraph_format.right_indent
                else 0
            )
            par_info['space_before'] = (
                int(p.paragraph_format.space_before.pt)
                if p.paragraph_format.space_before
                else int(ps.paragraph_format.space_before.pt)
                if ps.paragraph_format.space_before
                else int(psbs.paragraph_format.space_before.pt)
                if psbs and psbs.paragraph_format.space_before
                else 0
            )
            par_info['space_after'] = (
                int(p.paragraph_format.space_after.pt)
                if p.paragraph_format.space_after
                else int(ps.paragraph_format.space_after.pt)
                if ps.paragraph_format.space_after
                else int(psbs.paragraph_format.space_after.pt)
                if psbs and psbs.paragraph_format.space_after
                else 0
            )
            if 'Heading' in style:
                par_info['bold'] = (
                    p.runs[0].bold
                    if p.runs[0].bold
                    else ps.font.bold
                    if ps.font.bold
                    else psbs.font.bold
                    if psbs and psbs.font.bold
                    else False
                )
                par_info['italic'] = (
                    p.runs[0].italic
                    if p.runs[0].italic
                    else ps.font.italic
                    if ps.font.italic
                    else psbs.font.italic
                    if psbs and psbs.font.italic
                    else False
                )

            par_info['text'] = p.text

            if par_info['text'].startswith('Рис. '):
                style = 'ImageName'
            if par_info['text'].startswith('Таблица '):
                style = 'TableName'
            if paragraphs[i - 1].text.startswith('Таблица '):
                style = 'TableName'

            errors = detected_classes.get(
                style,
                RuleParams('Normal')
            ).validate(par_info)
            if errors:
                all_par_info.append({
                    'text': par_info['text'],
                    'style': style_info.get(style, 'Обычный'),
                    'errors': errors
                })
        else:
            if i != len(paragraphs) - 1:
                txt = paragraphs[i + 1].text
                if txt and txt.startswith('Рис. '):
                    style = 'ImageName'
                    par_info['line_spacing'] = 1.5
                    par_info['text'] = 'Является ' + txt
                else:
                    all_par_info.append({
                        'text': p.text,
                        'warning': ('Обнаружена пустая строка.'
                                    ' Возможна ошибка!'),
                        'next_text': txt
                    })
            else:
                all_par_info.append({
                    'text': p.text,
                    'warning': 'Обнаружена пустая строка. Возможна ошибка!',
                })
    return all_par_info


def get_table_info(doc: Document) -> dict:
    all_tab_info = []
    tbl = RuleParams('TableNormal')

    for ind, table in enumerate(doc.tables):
        table_info = dict()
        if not table.alignment or table.alignment.name != 'CENTER':
            table_info['table_alignment'] = (
                'Таблица должна располагаться по центру.'
            )
        table_info['rows'] = dict()
        for index, row in enumerate(table.rows):
            table_info['rows'][f'row_{index}'] = {}
            for index_1, cell in enumerate(row.cells):
                table_info['rows'][f'row_{index}'][f'cell_{index_1}'] = []
                for p in cell.paragraphs:
                    table_info_cell = dict()

                    if p.text and p.runs:
                        pr = p.runs[0]
                        ps = p.style
                        psbs = ps.base_style
                        table_info_cell['font_size'] = (
                            int(pr.font.size.pt)
                            if pr.font.size
                            else int(ps.font.size.pt)
                            if ps.font.size
                            else int(psbs.font.size.pt)
                            if psbs and psbs.font.size
                            else None
                        )
                        table_info_cell['font_name'] = (
                            pr.font.name
                            if pr.font.name
                            else ps.font.name
                            if ps.font.name
                            else psbs.font.name
                            if psbs and psbs.font.name
                            else None
                        )
                        table_info_cell['line_spacing'] = (
                            p.paragraph_format.line_spacing
                            if p.paragraph_format.line_spacing
                            else ps.paragraph_format.line_spacing
                            if ps.paragraph_format.line_spacing
                            else psbs.paragraph_format.line_spacing
                            if psbs and psbs.paragraph_format.line_spacing
                            else 1
                        )

                        table_info_cell['text'] = p.text
                        errors = tbl.validate(table_info_cell)

                        if errors:
                            table_info['rows'][f'row_{index}'][
                                f'cell_{index_1}'
                            ].append({'errors': errors,
                                      'text': table_info_cell['text']})
                if not table_info['rows'][f'row_{index}'][f'cell_{index_1}']:
                    table_info['rows'][f'row_{index}'].pop(f'cell_{index_1}')
            if not table_info['rows'][f'row_{index}']:
                table_info['rows'].pop(f'row_{index}')
        if table_info['rows']:
            all_tab_info.append((f'table_{ind}', table_info))
    raise
    return all_tab_info


def exam(doc: Document) -> tuple[list, list]:
    par_info = get_par_info(doc)
    tab_info = get_table_info(doc)
    return {'par_info': par_info, 'tab_info': tab_info}
